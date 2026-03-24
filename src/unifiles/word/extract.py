"""Word 文档内容提取功能。"""

from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE_TYPE

from ..exceptions import FileReadError
from ._table import list_to_markdown, table_to_list


def extract_text_docx(file_path: str) -> str:
    """提取 Word 文档完整文本，段落和表格按文档顺序输出。

    段落保留原始文本，表格转换为 Markdown 格式。

    Args:
        file_path: Word 文档路径

    Returns:
        文档完整文本，各部分之间用 ``\\n\\n`` 分隔

    Raises:
        FileNotFoundError: 文件不存在
        FileReadError: 读取文件时发生错误
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    try:
        document = Document(file_path)
        parts: list[str] = []

        for element in document.element.body:
            if element.tag.endswith("}p"):
                para = element.text
                if para and para.strip():
                    parts.append(para)
            elif element.tag.endswith("}tbl"):
                for table in document.tables:
                    if table._tbl is element:
                        data = table_to_list(table)
                        if data:
                            parts.append(list_to_markdown(data))
                        break

        return "\n\n".join(parts)
    except FileNotFoundError:
        raise
    except Exception as e:
        raise FileReadError(f"提取 Word 文本失败: {e}") from e


def extract_tables_docx(
    file_path: str,
    format: Literal["md", "list"] = "md",
) -> list[str] | list[list[list[str]]]:
    """提取 Word 文档中的所有表格。

    Args:
        file_path: Word 文档路径
        format: 输出格式，``"md"`` 返回 Markdown 字符串，``"list"`` 返回二维列表

    Returns:
        Markdown 格式列表或二维列表

    Raises:
        FileNotFoundError: 文件不存在
        FileReadError: 读取文件时发生错误
        ValueError: format 参数无效
    """
    if format not in ("md", "list"):
        raise ValueError(f"format 参数无效，期望 'md' 或 'list'，实际: {format}")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    try:
        document = Document(file_path)
        results: list[Any] = []

        for table in document.tables:
            data = table_to_list(table)
            if not data:
                continue
            if format == "list":
                results.append(data)
            else:
                results.append(list_to_markdown(data))

        return results
    except FileNotFoundError:
        raise
    except Exception as e:
        raise FileReadError(f"提取 Word 表格失败: {e}") from e


def extract_images_docx(
    file_path: str,
    output_dir: str | None = None,
) -> list[dict[str, Any]]:
    """提取 Word 文档中的图片并保存到指定目录。

    Args:
        file_path: Word 文档路径
        output_dir: 输出目录路径，默认为 ``./pics/``

    Returns:
        包含每张图片元信息的列表，字段包括 ``filename``、``path``、
        ``width``、``height``、``format``、``size_bytes``

    Raises:
        FileNotFoundError: 文件不存在
        FileReadError: 读取文件时发生错误
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    out_dir = Path(output_dir) if output_dir else Path("./pics/")

    try:
        document = Document(file_path)
        results: list[dict[str, Any]] = []
        filename_counts: dict[str, int] = {}

        for shape in document.inline_shapes:
            if shape.type != WD_INLINE_SHAPE_TYPE.PICTURE:
                continue

            r_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = document.part.related_parts[r_id]
            image_data = image_part.blob

            content_type = image_part.content_type
            ext = content_type.split("/")[-1].upper()
            format_map = {"JPEG": "JPG"}
            ext = format_map.get(ext, ext)

            base_name = f"image.{ext.lower()}"
            if base_name in filename_counts:
                filename_counts[base_name] += 1
                stem = Path(base_name).stem
                final_name = f"{stem}_{filename_counts[base_name]}.{ext.lower()}"
            else:
                filename_counts[base_name] = 0
                final_name = base_name

            out_dir.mkdir(parents=True, exist_ok=True)
            out_path = out_dir / final_name

            counter = 1
            while out_path.exists():
                stem = Path(base_name).stem
                final_name = f"{stem}_{counter}.{ext.lower()}"
                out_path = out_dir / final_name
                counter += 1

            out_path.write_bytes(image_data)

            results.append(
                {
                    "filename": final_name,
                    "path": str(out_path),
                    "width": shape.width,
                    "height": shape.height,
                    "format": ext,
                    "size_bytes": len(image_data),
                }
            )

        return results
    except FileNotFoundError:
        raise
    except Exception as e:
        raise FileReadError(f"提取 Word 图片失败: {e}") from e
