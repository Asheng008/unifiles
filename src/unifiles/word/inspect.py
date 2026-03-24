"""Word 文档综合检查功能。"""

from pathlib import Path
from typing import Any

from docx import Document

from ..exceptions import FileReadError
from ._table import list_to_markdown, table_to_list
from .extract import extract_images_docx


def inspect_docx(
    file_path: str,
    extract_images: bool = True,
    image_dir: str | None = None,
) -> dict[str, Any]:
    """综合检查 Word 文档元素。

    提取文档中的段落、表格和图片信息，返回综合统计结果。

    Args:
        file_path: Word 文档路径
        extract_images: 是否提取图片信息，默认 True
        image_dir: 图片输出目录路径，默认 ``./pics/``

    Returns:
        包含 ``paragraphs``、``tables``、``images``、``paragraph_count``、
        ``table_count``、``image_count`` 的字典

    Raises:
        FileNotFoundError: 文件不存在
        FileReadError: 读取文件时发生错误
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    try:
        document = Document(file_path)

        paragraphs = [para.text for para in document.paragraphs if para.text.strip()]

        tables = []
        for table in document.tables:
            data = table_to_list(table)
            if data:
                tables.append(list_to_markdown(data))

        images: list[dict[str, Any]] = []
        if extract_images:
            images = extract_images_docx(file_path, image_dir)

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "images": images,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "image_count": len(images),
        }
    except FileNotFoundError:
        raise
    except Exception as e:
        raise FileReadError(f"检查 Word 文档失败: {e}") from e
