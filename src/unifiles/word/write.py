"""Word 文档写入功能。"""

from pathlib import Path

from docx import Document

from ..exceptions import FileWriteError


def write_docx(content: str, file_path: str, title: str | None = None) -> None:
    """将内容写入 Word 文档。

    创建新的 Word 文档并写入内容。如果提供了标题，会将标题作为文档标题添加。

    Args:
        content: 要写入的文本内容
        file_path: 输出 Word 文档路径
        title: 可选的文档标题

    Raises:
        ValueError: 内容格式无效
        PermissionError: 文件权限不足
        FileWriteError: 写入文件时发生错误

    Example:
        >>> write_docx("Hello World", "output.docx")
        >>> write_docx("This is the content.", "output.docx", title="My Document")
    """
    if not isinstance(content, str):
        raise ValueError(f"内容格式无效，期望 str，实际类型: {type(content)}")

    try:
        document = Document()

        if title is not None:
            document.add_heading(title, level=0)

        lines = content.split("\n")
        for line in lines:
            document.add_paragraph(line)

        document.save(file_path)
    except PermissionError:
        raise
    except Exception as e:
        raise FileWriteError(f"写入 Word 文档失败: {e}") from e
