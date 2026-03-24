"""extract_images_docx 测试。"""

import os
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from unifiles.word import extract_images_docx


def test_extract_images_docx_success(tmp_path: Path, minimal_png: bytes):
    """测试成功提取图片。"""
    test_file = tmp_path / "test.docx"
    output_dir = tmp_path / "images"

    doc = Document()
    doc.add_paragraph("带图片的文档")
    doc.add_picture(BytesIO(minimal_png))
    doc.save(test_file)

    result = extract_images_docx(str(test_file), str(output_dir))
    assert isinstance(result, list)
    assert len(result) == 1
    img = result[0]
    assert img["filename"].endswith(".png")
    assert Path(img["path"]).exists()
    assert img["format"] == "PNG"
    assert img["size_bytes"] > 0
    assert img["width"] > 0
    assert img["height"] > 0


def test_extract_images_docx_no_images(tmp_path: Path):
    """测试无图片文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("只有文字")
    doc.save(test_file)

    result = extract_images_docx(str(test_file), str(tmp_path / "images"))
    assert result == []


def test_extract_images_docx_multiple_images(tmp_path: Path, minimal_png: bytes):
    """测试多图片提取。"""
    test_file = tmp_path / "test.docx"
    output_dir = tmp_path / "images"

    doc = Document()
    doc.add_picture(BytesIO(minimal_png))
    doc.add_picture(BytesIO(minimal_png))
    doc.save(test_file)

    result = extract_images_docx(str(test_file), str(output_dir))
    assert len(result) == 2
    filenames = [img["filename"] for img in result]
    assert len(set(filenames)) == 2


def test_extract_images_docx_default_output_dir(tmp_path: Path, minimal_png: bytes):
    """测试默认输出目录。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_picture(BytesIO(minimal_png))
    doc.save(test_file)

    original_cwd = os.getcwd()
    try:
        os.chdir(tmp_path)
        result = extract_images_docx(str(test_file))
        assert len(result) == 1
        assert Path(result[0]["path"]).exists()
    finally:
        os.chdir(original_cwd)


def test_extract_images_docx_file_not_found():
    """测试文件不存在。"""
    with pytest.raises(FileNotFoundError, match="文件不存在"):
        extract_images_docx("nonexistent.docx")
