"""extract_tables_docx 测试。"""

from pathlib import Path

import pytest
from docx import Document

from unifiles.word import extract_tables_docx


def test_extract_tables_docx_success(tmp_path: Path):
    """测试成功提取表格。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("文档标题")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "年龄"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "25"
    doc.save(test_file)

    result = extract_tables_docx(str(test_file))
    assert len(result) == 1
    assert "| 姓名 | 年龄 |" in result[0]
    assert "| 张三 | 25 |" in result[0]


def test_extract_tables_docx_no_tables(tmp_path: Path):
    """测试无表格文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("只有文字")
    doc.save(test_file)

    result = extract_tables_docx(str(test_file))
    assert result == []


def test_extract_tables_docx_empty_table(tmp_path: Path):
    """测试空表格。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_table(rows=1, cols=2)
    doc.save(test_file)

    result = extract_tables_docx(str(test_file))
    assert len(result) == 1
    assert "|  |  |" in result[0]


def test_extract_tables_docx_multiple_tables(tmp_path: Path):
    """测试多表格文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()

    table1 = doc.add_table(rows=2, cols=2)
    table1.cell(0, 0).text = "A"
    table1.cell(0, 1).text = "B"
    table1.cell(1, 0).text = "C"
    table1.cell(1, 1).text = "D"

    doc.add_paragraph("中间文字")

    table2 = doc.add_table(rows=2, cols=3)
    table2.cell(0, 0).text = "X"
    table2.cell(0, 1).text = "Y"
    table2.cell(0, 2).text = "Z"
    table2.cell(1, 0).text = "1"
    table2.cell(1, 1).text = "2"
    table2.cell(1, 2).text = "3"

    doc.save(test_file)

    result = extract_tables_docx(str(test_file))
    assert len(result) == 2
    assert "| A | B |" in result[0]
    assert "| X | Y | Z |" in result[1]


def test_extract_tables_docx_file_not_found():
    """测试文件不存在。"""
    with pytest.raises(FileNotFoundError, match="文件不存在"):
        extract_tables_docx("nonexistent.docx")


def test_extract_tables_docx_format_list(tmp_path: Path):
    """测试 format='list' 输出二维列表。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "年龄"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "25"
    doc.save(test_file)

    result = extract_tables_docx(str(test_file), format="list")
    assert len(result) == 1
    assert result[0] == [["姓名", "年龄"], ["张三", "25"]]


def test_extract_tables_docx_invalid_format(tmp_path: Path):
    """测试无效 format 参数。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.save(test_file)

    with pytest.raises(ValueError, match="format 参数无效"):
        extract_tables_docx(str(test_file), format="invalid")
