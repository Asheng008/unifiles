"""extract_text_docx 测试。"""

from pathlib import Path

import pytest
from docx import Document

from unifiles.word import extract_text_docx


def test_extract_text_docx_success(tmp_path: Path):
    """测试提取完整文本（含表格）。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("开头文字")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    doc.add_paragraph("结尾文字")
    doc.save(test_file)

    result = extract_text_docx(str(test_file))
    assert "开头文字" in result
    assert "结尾文字" in result
    assert "| A | B |" in result
    assert "| C | D |" in result


def test_extract_text_docx_no_tables(tmp_path: Path):
    """测试无表格文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段")
    doc.save(test_file)

    result = extract_text_docx(str(test_file))
    assert "第一段" in result
    assert "第二段" in result
    assert "|" not in result


def test_extract_text_docx_only_tables(tmp_path: Path):
    """测试只有表格的文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "X"
    table.cell(0, 1).text = "Y"
    table.cell(1, 0).text = "Z"
    table.cell(1, 1).text = "W"
    doc.save(test_file)

    result = extract_text_docx(str(test_file))
    assert "| X | Y |" in result
    assert "| Z | W |" in result


def test_extract_text_docx_file_not_found():
    """测试文件不存在。"""
    with pytest.raises(FileNotFoundError, match="文件不存在"):
        extract_text_docx("nonexistent.docx")
