"""inspect_docx 测试。"""

from pathlib import Path

import pytest
from docx import Document

from unifiles.word import inspect_docx


def test_inspect_docx_success(tmp_path: Path):
    """测试综合检查文档元素。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "X"
    table.cell(0, 1).text = "Y"
    table.cell(1, 0).text = "Z"
    table.cell(1, 1).text = "W"
    doc.save(test_file)

    result = inspect_docx(str(test_file), extract_images=False)
    assert result["paragraph_count"] == 2
    assert result["table_count"] == 1
    assert len(result["paragraphs"]) == 2
    assert len(result["tables"]) == 1
    assert "| X | Y |" in result["tables"][0]
    assert result["image_count"] == 0


def test_inspect_docx_empty(tmp_path: Path):
    """测试空文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.save(test_file)

    result = inspect_docx(str(test_file), extract_images=False)
    assert result["paragraph_count"] == 0
    assert result["table_count"] == 0
    assert result["image_count"] == 0


def test_inspect_docx_file_not_found():
    """测试文件不存在。"""
    with pytest.raises(FileNotFoundError, match="文件不存在"):
        inspect_docx("nonexistent.docx")
