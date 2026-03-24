"""read_docx 和 write_docx 测试。"""

from pathlib import Path

import pytest
from docx import Document

from unifiles.word import read_docx, write_docx
from unifiles.exceptions import FileReadError


def test_read_docx_success(tmp_path: Path):
    """测试成功读取 Word 文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("这是第一段。")
    doc.add_paragraph("这是第二段。")
    doc.add_paragraph("这是第三段。")
    doc.save(test_file)

    result = read_docx(str(test_file))
    assert isinstance(result, str)
    assert "第一段" in result
    assert "第二段" in result
    assert "第三段" in result
    lines = result.split("\n")
    assert len(lines) == 3


def test_read_docx_empty_paragraphs(tmp_path: Path):
    """测试读取包含空段落的文档。"""
    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("")  # 空段落
    doc.add_paragraph("第二段")
    doc.save(test_file)

    result = read_docx(str(test_file))
    lines = result.split("\n")
    assert len(lines) == 2
    assert "第一段" in result
    assert "第二段" in result


def test_read_docx_file_not_found():
    """测试文件不存在的情况。"""
    with pytest.raises(FileNotFoundError, match="文件不存在"):
        read_docx("nonexistent.docx")


def test_read_docx_invalid_format(tmp_path: Path):
    """测试无效格式文件。"""
    test_file = tmp_path / "test.txt"
    test_file.write_text("这不是 Word 文档")

    with pytest.raises(FileReadError):
        read_docx(str(test_file))


def test_write_docx_success(tmp_path: Path):
    """测试正常写入。"""
    test_file = tmp_path / "output.docx"
    content = "这是第一行\n这是第二行\n这是第三行"

    write_docx(content, str(test_file))

    assert test_file.exists()
    result = read_docx(str(test_file))
    assert "第一行" in result
    assert "第二行" in result
    assert "第三行" in result


def test_write_docx_with_title(tmp_path: Path):
    """测试带标题写入。"""
    test_file = tmp_path / "output.docx"
    content = "这是文档内容。"
    title = "我的文档标题"

    write_docx(content, str(test_file), title=title)

    assert test_file.exists()
    result = read_docx(str(test_file))
    assert title in result or "我的文档" in result
    assert "文档内容" in result


def test_write_docx_without_title(tmp_path: Path):
    """测试不带标题写入。"""
    test_file = tmp_path / "output.docx"
    content = "这是内容。"

    write_docx(content, str(test_file))

    assert test_file.exists()
    result = read_docx(str(test_file))
    assert "内容" in result


def test_write_docx_empty_content(tmp_path: Path):
    """测试写入空内容。"""
    test_file = tmp_path / "output.docx"
    content = ""

    write_docx(content, str(test_file))

    assert test_file.exists()
    result = read_docx(str(test_file))
    assert isinstance(result, str)


def test_write_docx_invalid_content(tmp_path: Path):
    """测试无效内容格式。"""
    test_file = tmp_path / "output.docx"

    with pytest.raises(ValueError, match="内容格式无效"):
        write_docx(123, str(test_file))  # type: ignore

    with pytest.raises(ValueError, match="内容格式无效"):
        write_docx(None, str(test_file))  # type: ignore


def test_write_docx_multiline_content(tmp_path: Path):
    """测试多行内容写入。"""
    test_file = tmp_path / "output.docx"
    content = "第一行\n\n第二行\n第三行"

    write_docx(content, str(test_file))

    result = read_docx(str(test_file))
    lines = result.split("\n")
    assert len(lines) >= 3
